from docx import Document

doc = Document()
doc.add_heading('John Doe - Software Engineer', 0)

doc.add_heading('Experience', level=1)
doc.add_paragraph('Developed web applications using HTML, CSS, JavaScript, and React. Built backend APIs using Node.js and Express.')

doc.add_heading('Skills', level=1)
doc.add_paragraph('Languages: JavaScript, TypeScript, Python')
doc.add_paragraph('Tools: Git, Docker, MongoDB')
doc.add_paragraph('Frameworks: React, Express, Node.js')

doc.save('dummy_resume.docx')
print("dummy_resume.docx created successfully!")
