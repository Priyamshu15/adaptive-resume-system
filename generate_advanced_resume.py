from docx import Document

doc = Document()
doc.add_heading('Jane Doe - AI Engineer', 0)

doc.add_heading('Professional Experience', level=1)

doc.add_heading('Senior Data Specialist | TechCorp Inc.', level=2)
doc.add_paragraph('June 2021 - Present')
# Strong bullets (with numbers/metrics)
doc.add_paragraph('- Developed machine learning models using Python and PyTorch that increased prediction accuracy by 25%.', style='List Bullet')
doc.add_paragraph('- Optimized SQL queries, reducing database load times by 40% across 3 major applications.', style='List Bullet')
# Weak bullets (no metrics)
doc.add_paragraph('- Worked on natural language processing tasks and neural networks.', style='List Bullet')
doc.add_paragraph('- Fixed bugs in the backend system and wrote documentation.', style='List Bullet')

doc.add_heading('Education', level=1)
doc.add_paragraph('B.S. in Computer Science - University of Technology (2020)')

doc.add_heading('Technical Skills', level=1)
doc.add_paragraph('Programming: Python, C++, Java')
doc.add_paragraph('Machine Learning: PyTorch, Neural Networks, Computer Vision')
doc.add_paragraph('Other: SQL, Docker, Linux')

doc.save('advanced_dummy_resume.docx')
print("advanced_dummy_resume.docx created successfully!")
